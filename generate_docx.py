import os
from docx import Document

output_dir = r"C:\Users\Thinh Phat\Documents\template"
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

doc = Document()
doc.add_heading('Comprehensive Test Case Suite for SE114 Calling App', 0)

doc.add_paragraph("This test suite covers the primary functional modules of the SE114 Calling App based on its architecture (Authentication, Profile & Friends, Chat System, Posts & Comments, Server Management, and Calling System).")

def add_module_table(module_name, rows):
    doc.add_heading(module_name, level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Test Case ID', 'Scenario', 'Pre-conditions', 'Steps', 'Expected Result']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
    
    for r in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(r):
            row_cells[i].text = val
    doc.add_paragraph()

# 1. Authentication
rows_auth = [
    ["TC_AUTH_01", "Successful Registration", "User is on the Register screen", "1. Enter valid email, username, and password.\n2. Click 'Register'.", "Account is created successfully and user is redirected to the Login or Home screen."],
    ["TC_AUTH_02", "Registration with existing email", "User is on the Register screen", "1. Enter an email that is already registered.\n2. Click 'Register'.", "Error message is displayed: \"Email already in use\"."],
    ["TC_AUTH_03", "Successful Login", "User is on the Login screen", "1. Enter valid credentials.\n2. Click 'Login'.", "User is logged in and redirected to HomePageActivity."],
    ["TC_AUTH_04", "Login with invalid credentials", "User is on the Login screen", "1. Enter incorrect password or non-existent email.\n2. Click 'Login'.", "Error message is displayed: \"Invalid email or password\"."],
    ["TC_AUTH_05", "Logout functionality", "User is logged in and on HomePage", "1. Open Profile/Settings.\n2. Click 'Logout'.", "User session ends and is redirected to the Login screen."]
]
add_module_table("1. Module: Authentication (Login & Registration)", rows_auth)

# 2. Profile & Friends
rows_prof = [
    ["TC_PROF_01", "View User Profile", "User is logged in", "1. Navigate to ProfileActivity.", "User's avatar, username, and status are displayed correctly."],
    ["TC_PROF_02", "Edit Profile Information", "User is on Profile screen", "1. Click 'Edit Profile'.\n2. Change username and avatar.\n3. Save.", "Profile is updated and changes reflect across the app immediately."],
    ["TC_FRND_01", "Send Friend Request", "User is logged in", "1. Go to ManageFriendsActivity.\n2. Search for a user.\n3. Click 'Add Friend'.", "Friend request is sent. Target user receives a notification."],
    ["TC_FRND_02", "Accept/Reject Request", "User has a pending request", "1. Go to pending requests.\n2. Click 'Accept' or 'Reject'.", "If accepted, user is added to friends list. If rejected, request is removed."],
    ["TC_FRND_03", "Remove Friend", "User is on Friends List", "1. Select a friend.\n2. Choose 'Remove Friend'.", "User is removed from the friends list."]
]
add_module_table("2. Module: Profile & Friends Management", rows_prof)

# 3. Chat System
rows_chat = [
    ["TC_CHAT_01", "Send Text Message", "User is in a ChatDetailActivity", "1. Type a message.\n2. Click Send.", "Message appears in the chat bubble immediately with current timestamp."],
    ["TC_CHAT_02", "Send Media/Document", "User is in a ChatDetailActivity", "1. Click Attachment icon.\n2. Select an image/document.\n3. Send.", "File uploads and appears in the chat. Viewers (Image/Document Viewer) can open it."],
    ["TC_CHAT_03", "Background Notifications", "App is in background", "1. Another user sends a message to the current user.", "User receives a push notification via MessageNotificationService."],
    ["TC_CHAT_04", "Search in Chat", "User is in a ChatDetailActivity", "1. Open SearchInChatActivity.\n2. Type a keyword.", "Matching messages in the current conversation are displayed."],
    ["TC_CHAT_05", "Mention User", "User is in a Server/Group chat", "1. Type '@' and select a user from MentionAdapter list.", "The mentioned user is highlighted and receives a specific mention notification."]
]
add_module_table("3. Module: Chat System", rows_chat)

# 4. Posts & Comments
rows_post = [
    ["TC_POST_01", "Create a New Post", "User is in PostChannelActivity", "1. Click 'Create Post'.\n2. Add title, content, and image.\n3. Publish.", "Post is published and appears in the channel feed."],
    ["TC_POST_02", "View Post Channel", "User navigates to a channel", "1. Open PostChannelActivity.", "List of posts loads correctly, displaying author, timestamp, and content snippet."],
    ["TC_POST_03", "Add Comment", "User is viewing a Post", "1. Scroll to comment section.\n2. Type comment.\n3. Submit.", "Comment is added to the post via CommentAdapter."],
    ["TC_POST_04", "View Comments", "User opens PostCommentActivity", "1. Click on a post to view comments.", "All previous comments load sequentially."]
]
add_module_table("4. Module: Posts & Comments", rows_post)

# 5. Server Management
rows_serv = [
    ["TC_SERV_01", "View Server Channels", "User is in ServerViewerActivity", "1. Select a server from the list.", "The layout updates to show available text, voice, and post channels for that server."],
    ["TC_SERV_02", "Manage Members", "User is Server Admin", "1. Open ManageMembersActivity.\n2. Select a member.\n3. Kick/Ban.", "Selected member is removed from the server and loses access."],
    ["TC_SERV_03", "Change Theme/Color", "User is in Chat/Server Settings", "1. Open ChangeColorActivity.\n2. Select a new color theme.", "The chat/server UI updates to reflect the new color choices."]
]
add_module_table("5. Module: Server / Channel Management", rows_serv)

# 6. Calling & Screen Sharing
rows_call = [
    ["TC_CALL_01", "Initiate Voice/Video Call", "User is in ChatDetailActivity", "1. Click 'Call' or 'Video Call' button.", "CallDetailActivity opens and starts ringing the recipient."],
    ["TC_CALL_02", "Receive Call", "User is logged in", "1. Incoming call triggered.", "Call incoming UI is shown. User can Accept or Decline."],
    ["TC_CALL_03", "Mute/Unmute Audio", "User is in an active call", "1. Click the Microphone icon.", "Audio stream is muted/unmuted for other participants."],
    ["TC_CALL_04", "Enable/Disable Camera", "User is in an active video call", "1. Click the Camera icon.", "Video feed toggles on/off."],
    ["TC_CALL_05", "Screen Sharing", "User is in an active call", "1. Click 'Share Screen'.\n2. Grant Android projection permission.", "MyScreenShareService starts; other participants can see the user's screen."],
    ["TC_CALL_06", "End Call", "User is in an active call", "1. Click 'End Call' button.", "Call terminates, and user is returned to the previous screen."]
]
add_module_table("6. Module: Calling & Screen Sharing", rows_call)

output_path = os.path.join(output_dir, "SE114_CallingApp_Test_Cases.docx")
doc.save(output_path)
print("Docx created at: " + output_path)
